"""
Generate sample vendor contracts for testing the extraction automation
Creates realistic Word documents with contract information
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path


def create_sample_contracts():
    """Create 3 sample vendor contracts in Word format"""

    sample_contracts_dir = Path("sample-contracts")
    if not sample_contracts_dir.exists():
        sample_contracts_dir.mkdir(exist_ok=True)

    # Sample Contract 1: IT Services
    doc1 = Document()
    doc1.add_heading("VENDOR SERVICE AGREEMENT", level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc1.add_paragraph("").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc1.add_heading("CONTRACT HEADER", level=1)
    doc1.add_paragraph(f"Contract Number: VC-2025-0001")
    doc1.add_paragraph(f"Vendor Name: TechSolutions Inc.")
    doc1.add_paragraph(f"Contract Start Date: January 1, 2025")
    doc1.add_paragraph(f"Contract End Date: December 31, 2025")
    doc1.add_paragraph(f"Contract Value: $150,000 USD")
    doc1.add_paragraph(f"Payment Terms: Net 30 days")
    doc1.add_paragraph(f"Currency: USD")
    doc1.add_paragraph(f"Contract Type: Service Agreement")

    doc1.add_heading("SERVICES & RATES", level=1)

    doc1.add_heading("Service 1: Cloud Infrastructure Support", level=2)
    doc1.add_paragraph(f"Service Name: Cloud Infrastructure Support")
    doc1.add_paragraph(f"Description: 24/7 monitoring, maintenance, and technical support for cloud infrastructure including AWS, Azure, and GCP services.")
    doc1.add_paragraph(f"Unit: Per Month")
    doc1.add_paragraph(f"Rate: $5,000 per month")
    doc1.add_paragraph(f"Currency: USD")
    doc1.add_paragraph(f"Minimum Order: 1 month")
    doc1.add_paragraph(f"Volume Discount: 10% discount for annual commitment")
    doc1.add_paragraph(f"Effective From: January 1, 2025")

    doc1.add_heading("Service 2: Security Audit & Penetration Testing", level=2)
    doc1.add_paragraph(f"Service Name: Security Audit & Penetration Testing")
    doc1.add_paragraph(f"Description: Quarterly comprehensive security audits, vulnerability assessments, and authorized penetration testing of infrastructure and applications.")
    doc1.add_paragraph(f"Unit: Per Quarter")
    doc1.add_paragraph(f"Rate: $8,000 per quarter")
    doc1.add_paragraph(f"Currency: USD")
    doc1.add_paragraph(f"Minimum Order: 1 quarter (3 months)")
    doc1.add_paragraph(f"Volume Discount: 15% discount for 4-quarter commitment")
    doc1.add_paragraph(f"Effective From: January 1, 2025")

    doc1.add_heading("Service 3: On-Site Technical Support", level=2)
    doc1.add_paragraph(f"Service Name: On-Site Technical Support")
    doc1.add_paragraph(f"Description: On-site technical support for critical infrastructure issues and emergency troubleshooting.")
    doc1.add_paragraph(f"Unit: Per Hour")
    doc1.add_paragraph(f"Rate: $150 per hour")
    doc1.add_paragraph(f"Currency: USD")
    doc1.add_paragraph(f"Minimum Order: 4 hours per engagement")
    doc1.add_paragraph(f"Volume Discount: 10% discount for blocks of 100+ hours")
    doc1.add_paragraph(f"Effective From: January 1, 2025")

    contract1_path = sample_contracts_dir / "sample_contract_1_techsolutions.docx"
    doc1.save(contract1_path)
    print(f"✓ Created: {contract1_path}")

    # Sample Contract 2: Office Supplies
    doc2 = Document()
    doc2.add_heading("PROCUREMENT AGREEMENT", level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc2.add_paragraph("")

    doc2.add_heading("CONTRACT DETAILS", level=1)
    doc2.add_paragraph(f"Contract Number: VC-2025-0002")
    doc2.add_paragraph(f"Vendor Name: OfficeWorld Supplies Ltd.")
    doc2.add_paragraph(f"Contract Start Date: February 1, 2025")
    doc2.add_paragraph(f"Contract End Date: January 31, 2026")
    doc2.add_paragraph(f"Contract Value: $45,000 USD")
    doc2.add_paragraph(f"Payment Terms: Net 15 days")
    doc2.add_paragraph(f"Currency: USD")
    doc2.add_paragraph(f"Contract Type: Supply Agreement")

    doc2.add_heading("PRODUCTS & PRICING", level=1)

    doc2.add_heading("Product 1: Premium Printer Paper", level=2)
    doc2.add_paragraph(f"Service Name: Premium Printer Paper A4 80gsm")
    doc2.add_paragraph(f"Description: High quality white printer paper, A4 size, 80 gsm weight. Sold in reams of 500 sheets.")
    doc2.add_paragraph(f"Unit: Per Ream")
    doc2.add_paragraph(f"Rate: $5.50 per ream")
    doc2.add_paragraph(f"Currency: USD")
    doc2.add_paragraph(f"Minimum Order: 10 reams")
    doc2.add_paragraph(f"Volume Discount: 5% for 100+ reams, 10% for 500+ reams per order")
    doc2.add_paragraph(f"Effective From: February 1, 2025")

    doc2.add_heading("Product 2: Ballpoint Pens", level=2)
    doc2.add_paragraph(f"Service Name: Ballpoint Pens - Black/Blue/Red")
    doc2.add_paragraph(f"Description: Professional ballpoint pens available in black, blue, and red. Sold in boxes of 50 pens.")
    doc2.add_paragraph(f"Unit: Per Box")
    doc2.add_paragraph(f"Rate: $12.00 per box")
    doc2.add_paragraph(f"Currency: USD")
    doc2.add_paragraph(f"Minimum Order: 5 boxes")
    doc2.add_paragraph(f"Volume Discount: 8% for monthly orders of 200+ boxes")
    doc2.add_paragraph(f"Effective From: February 1, 2025")

    doc2.add_heading("Product 3: Desk Organizers", level=2)
    doc2.add_paragraph(f"Service Name: Desk Organizer Sets")
    doc2.add_paragraph(f"Description: Modular desk organizer with compartments for pens, papers, and office supplies.")
    doc2.add_paragraph(f"Unit: Per Unit")
    doc2.add_paragraph(f"Rate: $25.00 per unit")
    doc2.add_paragraph(f"Currency: USD")
    doc2.add_paragraph(f"Minimum Order: 10 units")
    doc2.add_paragraph(f"Volume Discount: 12% for bulk orders of 500+ units")
    doc2.add_paragraph(f"Effective From: February 1, 2025")

    contract2_path = sample_contracts_dir / "sample_contract_2_officeworld.docx"
    doc2.save(contract2_path)
    print(f"✓ Created: {contract2_path}")

    # Sample Contract 3: Logistics Services
    doc3 = Document()
    doc3.add_heading("LOGISTICS SERVICES AGREEMENT", level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc3.add_paragraph("")

    doc3.add_heading("AGREEMENT DETAILS", level=1)
    doc3.add_paragraph(f"Contract Number: VC-2025-0003")
    doc3.add_paragraph(f"Vendor Name: GlobalShip Logistics")
    doc3.add_paragraph(f"Contract Start Date: March 15, 2025")
    doc3.add_paragraph(f"Contract End Date: March 14, 2026")
    doc3.add_paragraph(f"Contract Value: $250,000 USD")
    doc3.add_paragraph(f"Payment Terms: Net 45 days")
    doc3.add_paragraph(f"Currency: USD")
    doc3.add_paragraph(f"Contract Type: Logistics Services Agreement")

    doc3.add_heading("SERVICES & RATES", level=1)

    doc3.add_heading("Service 1: Domestic Shipping", level=2)
    doc3.add_paragraph(f"Service Name: Domestic Ground Shipping")
    doc3.add_paragraph(f"Description: Standard ground shipping for packages within continental United States. Delivery within 5-7 business days.")
    doc3.add_paragraph(f"Unit: Per Pound")
    doc3.add_paragraph(f"Rate: $0.85 per pound")
    doc3.add_paragraph(f"Currency: USD")
    doc3.add_paragraph(f"Minimum Order: 100 pounds per shipment")
    doc3.add_paragraph(f"Volume Discount: 10% for monthly volumes exceeding 10,000 pounds")
    doc3.add_paragraph(f"Effective From: March 15, 2025")

    doc3.add_heading("Service 2: International Shipping", level=2)
    doc3.add_paragraph(f"Service Name: International Express Shipping")
    doc3.add_paragraph(f"Description: Express international shipping to Europe, Asia, and other regions. Delivery within 3-5 business days.")
    doc3.add_paragraph(f"Unit: Per Pound")
    doc3.add_paragraph(f"Rate: $2.50 per pound")
    doc3.add_paragraph(f"Currency: USD")
    doc3.add_paragraph(f"Minimum Order: 50 pounds per shipment")
    doc3.add_paragraph(f"Volume Discount: 15% for monthly international volumes over 5,000 pounds")
    doc3.add_paragraph(f"Effective From: March 15, 2025")

    doc3.add_heading("Service 3: Warehouse Storage", level=2)
    doc3.add_paragraph(f"Service Name: Warehouse Storage & Fulfillment")
    doc3.add_paragraph(f"Description: Climate-controlled warehouse storage with inventory management and order fulfillment services.")
    doc3.add_paragraph(f"Unit: Per Pallet per Month")
    doc3.add_paragraph(f"Rate: $75 per pallet per month")
    doc3.add_paragraph(f"Currency: USD")
    doc3.add_paragraph(f"Minimum Order: 10 pallets per month")
    doc3.add_paragraph(f"Volume Discount: 20% discount for 6-month commitment of 50+ pallets")
    doc3.add_paragraph(f"Effective From: March 15, 2025")

    contract3_path = sample_contracts_dir / "sample_contract_3_globalship.docx"
    doc3.save(contract3_path)
    print(f"✓ Created: {contract3_path}")

    print(f"\n✓ Sample contracts created in {sample_contracts_dir}/ directory")
    print(f"  - sample_contract_1_techsolutions.docx")
    print(f"  - sample_contract_2_officeworld.docx")
    print(f"  - sample_contract_3_globalship.docx")


if __name__ == "__main__":
    print("Creating Sample Vendor Contracts")
    print("=" * 50)
    create_sample_contracts()
    print("\nYou can now run: python contract_extractor.py")
    print("to extract information from these sample contracts!")
